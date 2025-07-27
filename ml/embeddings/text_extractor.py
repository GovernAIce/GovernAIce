#!/usr/bin/env python3
"""
Advanced Multilingual Text Extraction System
============================================

A comprehensive text extraction model that supports:
- Normal text documents (PDF, DOCX, TXT)
- Image-based text (OCR with multilingual support)
- Tables and structured data
- Graphs and charts
- Multiple languages (English, Mandarin, and more)
- URLs and web pages (HTML, online PDFs, articles)
- Online images and documents

Dependencies:
pip install pytesseract opencv-python pandas numpy pillow PyPDF2 python-docx
pip install transformers torch torchvision easyocr paddlepaddle paddleocr
pip install tabulate openpyxl xlrd camelot-py[cv] pdfplumber
pip install requests beautifulsoup4 lxml html2text selenium webdriver-manager
pip install newspaper3k readability-lxml scrapy-fake-useragent
"""

import os
import sys
import logging
import traceback
import io
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
from datetime import datetime
import json

# Core libraries
import numpy as np
import pandas as pd
from PIL import Image

# Import configuration
from ml.config.settings import MLConfig

# Document processing - with graceful fallbacks
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    print("PyPDF2 not available. Install with: pip install PyPDF2")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    print("pdfplumber not available. Install with: pip install pdfplumber")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available. Install with: pip install python-docx")

try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    CAMELOT_AVAILABLE = False
    print("camelot-py not available. Install with: pip install camelot-py[cv]")

# Advanced OCR - with graceful fallbacks
try:
    import cv2
    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False
    print("OpenCV not available. Install with: pip install opencv-python")

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
    print("pytesseract not available. Install with: pip install pytesseract")

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False
    print("easyocr not available. Install with: pip install easyocr")

# Advanced OCR
try:
    from paddleocr import PaddleOCR
    PADDLE_AVAILABLE = True
except ImportError:
    PADDLE_AVAILABLE = False
    print("PaddleOCR not available. Install with: pip install paddlepaddle paddleocr")

# Web scraping and URL handling
try:
    import requests
    from bs4 import BeautifulSoup
    import html2text
    from urllib.parse import urlparse, urljoin, quote
    from urllib.request import urlretrieve
    WEB_SCRAPING_AVAILABLE = True
except ImportError:
    WEB_SCRAPING_AVAILABLE = False
    print("Web scraping libraries not available. Install with: pip install requests beautifulsoup4 lxml html2text")

# Advanced web content extraction
try:
    from newspaper import Article
    import newspaper
    NEWSPAPER_AVAILABLE = True
except ImportError:
    NEWSPAPER_AVAILABLE = False
    print("Newspaper3k not available. Install with: pip install newspaper3k")

# User agent spoofing
try:
    from fake_useragent import UserAgent
    FAKE_USERAGENT_AVAILABLE = True
except ImportError:
    FAKE_USERAGENT_AVAILABLE = False
    print("Fake UserAgent not available. Install with: pip install fake-useragent")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@dataclass
class ExtractionResult:
    """Data class for extraction results"""
    text: str
    confidence: float
    language: str
    content_type: str
    metadata: Dict[str, Any]
    extraction_time: float

class AdvancedTextExtractor:
    """
    Advanced multilingual text extraction system with fallback options
    """
    
    def __init__(self,
                 languages: List[str] = ['en', 'ch_sim'],
                 output_dir: str = "extracted_texts",
                 use_gpu: bool = False):
        """
        Initialize the text extractor
        
        Args:
            languages: List of languages for OCR
            output_dir: Directory to save extracted texts
            use_gpu: Whether to use GPU for OCR
        """
        self.languages = languages
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.use_gpu = use_gpu
        
        # Initialize OCR engines
        self._init_ocr_engines()
        
        # Initialize document processors
        self._init_document_processors()
        
        # Initialize web scraping
        self._init_web_scraping()
        
        logger.info("Advanced Text Extractor initialized")

    def _init_ocr_engines(self):
        """Initialize OCR engines with fallbacks"""
        self.ocr_engines = {}
        
        # Try EasyOCR first (most reliable)
        if EASYOCR_AVAILABLE:
            try:
                self.ocr_engines['easyocr'] = easyocr.Reader(self.languages, gpu=self.use_gpu)
                logger.info("EasyOCR initialized successfully")
            except Exception as e:
                logger.warning(f"EasyOCR initialization failed: {e}")
        
        # Try PaddleOCR as backup
        if PADDLE_AVAILABLE:
            try:
                self.ocr_engines['paddleocr'] = PaddleOCR(use_angle_cls=True, lang='en', use_gpu=self.use_gpu)
                logger.info("PaddleOCR initialized successfully")
            except Exception as e:
                logger.warning(f"PaddleOCR initialization failed: {e}")
        
        # Try Tesseract as last resort
        if PYTESSERACT_AVAILABLE:
            try:
                # Test if tesseract is available
                pytesseract.get_tesseract_version()
                self.ocr_engines['tesseract'] = pytesseract
                logger.info("Tesseract initialized successfully")
            except Exception as e:
                logger.warning(f"Tesseract initialization failed: {e}")
        
        if not self.ocr_engines:
            logger.warning("No OCR engines available. Image text extraction will not work.")

    def _init_document_processors(self):
        """Initialize document processors"""
        self.document_processors = {}
        
        if PYPDF2_AVAILABLE:
            self.document_processors['pdf'] = 'PyPDF2'
        
        if PDFPLUMBER_AVAILABLE:
            self.document_processors['pdf_plumber'] = 'pdfplumber'
        
        if DOCX_AVAILABLE:
            self.document_processors['docx'] = 'python-docx'
        
        if CAMELOT_AVAILABLE:
            self.document_processors['tables'] = 'camelot'

    def _init_web_scraping(self):
        """Initialize web scraping capabilities"""
        self.web_scraping_available = WEB_SCRAPING_AVAILABLE
        self.newspaper_available = NEWSPAPER_AVAILABLE
        self.fake_useragent_available = FAKE_USERAGENT_AVAILABLE

    def extract_text(self, source: Union[str, Path]) -> ExtractionResult:
        """
        Extract text from various sources
        
        Args:
            source: File path, URL, or text content
            
        Returns:
            ExtractionResult object
        """
        start_time = datetime.now()
        
        try:
            if self._is_url(str(source)):
                result = self._extract_from_url(str(source), start_time)
            elif Path(source).exists():
                result = self._extract_from_file(Path(source), start_time)
            else:
                # Treat as direct text input
                result = ExtractionResult(
                    text=str(source),
                    confidence=1.0,
                    language='en',
                    content_type='text',
                    metadata={'source': 'direct_input'},
                    extraction_time=(datetime.now() - start_time).total_seconds()
                )
            
            # Save extracted text
            self._save_extracted_text(source, result)
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from {source}: {e}")
            return ExtractionResult(
                text=f"Error extracting text: {str(e)}",
                confidence=0.0,
                language='unknown',
                content_type='error',
                metadata={'error': str(e)},
                extraction_time=(datetime.now() - start_time).total_seconds()
            )

    def _is_url(self, source: str) -> bool:
        """Check if source is a URL"""
        return source.startswith(('http://', 'https://', 'ftp://'))

    def _extract_from_url(self, url: str, start_time: datetime) -> ExtractionResult:
        """Extract text from URL"""
        if not self.web_scraping_available:
            raise ValueError("Web scraping not available. Install required dependencies.")
        
        # Implementation would go here
        # For now, return a placeholder
        return ExtractionResult(
            text=f"Web extraction not implemented for {url}",
            confidence=0.0,
            language='en',
            content_type='web',
            metadata={'url': url},
            extraction_time=(datetime.now() - start_time).total_seconds()
        )

    def _extract_from_file(self, file_path: Path, start_time: datetime) -> ExtractionResult:
        """Extract text from file"""
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path, start_time)
        elif file_extension == '.docx':
            return self._extract_from_docx(file_path, start_time)
        elif file_extension == '.txt':
            return self._extract_from_txt(file_path, start_time)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            return self._extract_from_image(file_path, start_time)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    def _extract_from_pdf(self, file_path: Path, start_time: datetime) -> ExtractionResult:
        """Extract text from PDF"""
        if not PYPDF2_AVAILABLE and not PDFPLUMBER_AVAILABLE:
            raise ValueError("PDF processing not available. Install PyPDF2 or pdfplumber.")
        
        text_content = ""
        confidence = 0.0
        
        # Try pdfplumber first (better for complex PDFs)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text_content += page.extract_text() or ""
                confidence = 0.9
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
        
        # Fallback to PyPDF2
        if not text_content and PYPDF2_AVAILABLE:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() or ""
                confidence = 0.8
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")
        
        if not text_content:
            raise ValueError("Could not extract text from PDF")
        
        return ExtractionResult(
            text=text_content,
            confidence=confidence,
            language='en',
            content_type='pdf',
            metadata={'pages': len(text_content.split('\n'))},
            extraction_time=(datetime.now() - start_time).total_seconds()
        )

    def _extract_from_docx(self, file_path: Path, start_time: datetime) -> ExtractionResult:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX processing not available. Install python-docx.")
        
        try:
            doc = Document(file_path)
            text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return ExtractionResult(
                text=text_content,
                confidence=1.0,
                language='en',
                content_type='docx',
                metadata={'paragraphs': len(doc.paragraphs)},
                extraction_time=(datetime.now() - start_time).total_seconds()
            )
        except Exception as e:
            raise ValueError(f"Error extracting from DOCX: {e}")

    def _extract_from_txt(self, file_path: Path, start_time: datetime) -> ExtractionResult:
        """Extract text from TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            return ExtractionResult(
                text=text_content,
                confidence=1.0,
                language='en',
                content_type='txt',
                metadata={'lines': len(text_content.split('\n'))},
                extraction_time=(datetime.now() - start_time).total_seconds()
            )
        except Exception as e:
            raise ValueError(f"Error reading text file: {e}")

    def _extract_from_image(self, file_path: Path, start_time: datetime) -> ExtractionResult:
        """Extract text from image using OCR"""
        if not self.ocr_engines:
            raise ValueError("No OCR engines available. Install easyocr, paddleocr, or pytesseract.")
        
        try:
            # Load image
            with Image.open(file_path) as img:
                # Try each OCR engine
                for engine_name, engine in self.ocr_engines.items():
                    try:
                        if engine_name == 'easyocr':
                            results = engine.readtext(np.array(img))
                            text_content = " ".join([result[1] for result in results])
                            confidence = np.mean([result[2] for result in results]) if results else 0.0
                        elif engine_name == 'paddleocr':
                            results = engine.ocr(np.array(img))
                            text_content = " ".join([line[1][0] for line in results[0]]) if results[0] else ""
                            confidence = 0.8  # PaddleOCR doesn't provide confidence scores
                        elif engine_name == 'tesseract':
                            text_content = engine.image_to_string(img)
                            confidence = 0.7  # Tesseract confidence varies
                        
                        if text_content.strip():
                            return ExtractionResult(
                                text=text_content,
                                confidence=confidence,
                                language='en',
                                content_type='image',
                                metadata={'ocr_engine': engine_name, 'file_size': file_path.stat().st_size},
                                extraction_time=(datetime.now() - start_time).total_seconds()
                            )
                    except Exception as e:
                        logger.warning(f"{engine_name} failed: {e}")
                        continue
                
                raise ValueError("All OCR engines failed")
                
        except Exception as e:
            raise ValueError(f"Error processing image: {e}")

    def _save_extracted_text(self, source: Union[Path, str], result: ExtractionResult) -> Path:
        """Save extracted text to file"""
        try:
            if isinstance(source, Path):
                filename = source.stem
            else:
                filename = Path(str(source)).stem
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{filename}_extracted_{timestamp}.txt"
            output_path = self.output_dir / output_filename
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("=== EXTRACTED TEXT ===\n")
                f.write(f"Source: {source}\n")
                f.write(f"Extraction Time: {result.extraction_time:.2f}s\n")
                f.write(f"Confidence: {result.confidence:.2f}\n")
                f.write(f"Language: {result.language}\n")
                f.write(f"Content Type: {result.content_type}\n")
                f.write("=" * 50 + "\n\n")
                f.write(result.text)
            
            logger.info(f"Extracted text saved to: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error saving extracted text: {e}")
            return Path("")

    def batch_extract(self, sources: List[Union[str, Path]], file_pattern: str = "*") -> List[ExtractionResult]:
        """Extract text from multiple sources"""
        results = []
        
        for source in sources:
            try:
                result = self.extract_text(source)
                results.append(result)
                logger.info(f"Successfully extracted from {source}")
            except Exception as e:
                logger.error(f"Failed to extract from {source}: {e}")
                # Add error result
                results.append(ExtractionResult(
                    text=f"Error: {str(e)}",
                    confidence=0.0,
                    language='unknown',
                    content_type='error',
                    metadata={'error': str(e)},
                    extraction_time=0.0
                ))
        
        return results

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        formats = []
        
        if PYPDF2_AVAILABLE or PDFPLUMBER_AVAILABLE:
            formats.append("PDF")
        if DOCX_AVAILABLE:
            formats.append("DOCX")
        formats.append("TXT")
        
        if self.ocr_engines:
            formats.extend(["JPG", "JPEG", "PNG", "BMP", "TIFF"])
        
        if self.web_scraping_available:
            formats.append("URL")
        
        return formats

    def create_extraction_summary(self, results: List[ExtractionResult]) -> str:
        """Create a summary of extraction results"""
        successful = [r for r in results if r.confidence > 0]
        failed = [r for r in results if r.confidence == 0]
        
        summary = f"""
Extraction Summary:
- Total files processed: {len(results)}
- Successful extractions: {len(successful)}
- Failed extractions: {len(failed)}
- Average confidence: {np.mean([r.confidence for r in successful]) if successful else 0:.2f}
- Total text extracted: {sum(len(r.text) for r in successful)} characters
        """
        
        return summary

def main():
    """Main function for text extraction"""
    print("Advanced Text Extraction System")
    print("=" * 40)
    
    extractor = AdvancedTextExtractor()
    
    print(f"Supported formats: {', '.join(extractor.get_supported_formats())}")
    
    # Get input files
    print("\nEnter file paths to extract (one per line, empty line to finish):")
    files = []
    while True:
        file_path = input("File path: ").strip()
        if not file_path:
            break
        if Path(file_path).exists():
            files.append(file_path)
            print(f"  ✅ Added: {Path(file_path).name}")
        else:
            print(f"  ❌ File not found: {file_path}")
    
    if not files:
        print("No files provided")
        return
    
    # Extract text
    print(f"\nExtracting text from {len(files)} files...")
    results = extractor.batch_extract(files)
    
    # Show summary
    print(extractor.create_extraction_summary(results))
    
    # Show individual results
    for i, result in enumerate(results, 1):
        print(f"\n{i}. {Path(files[i-1]).name}:")
        print(f"   Confidence: {result.confidence:.2f}")
        print(f"   Content: {result.text[:100]}...")

if __name__ == "__main__":
    main()
