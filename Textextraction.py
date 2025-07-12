#!/usr/bin/env python3
"""
Advanced Multilingual Text Extraction System
============================================

A comprehensive text extraction model that supports:
- Normal text documents (PDF, DOCX, TXT)
- Image-based text (OCR with multilingual support)
- Tables and structured data
- Graphs and charts
- Multiple languages (English, Mandarin, and more)
- URLs and web pages (HTML, online PDFs, articles)
- Online images and documents

Dependencies:
pip install pytesseract opencv-python pandas numpy pillow PyPDF2 python-docx
pip install transformers torch torchvision easyocr paddlepaddle paddleocr
pip install tabulate openpyxl xlrd camelot-py[cv] pdfplumber
pip install requests beautifulsoup4 lxml html2text selenium webdriver-manager
pip install newspaper3k readability-lxml scrapy-fake-useragent
"""

import os
import sys
import logging
import traceback
import io
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
from datetime import datetime
import json

# Core libraries
import cv2
import numpy as np
import pandas as pd
from PIL import Image
import pytesseract
import easyocr

# Document processing
import PyPDF2
import pdfplumber
from docx import Document
import camelot

# Advanced OCR
try:
    from paddleocr import PaddleOCR
    PADDLE_AVAILABLE = True
except ImportError:
    PADDLE_AVAILABLE = False
    print("PaddleOCR not available. Install with: pip install paddlepaddle paddleocr")

# Web scraping and URL handling
try:
    import requests
    from bs4 import BeautifulSoup
    import html2text
    from urllib.parse import urlparse, urljoin, quote
    from urllib.request import urlretrieve
    WEB_SCRAPING_AVAILABLE = True
except ImportError:
    WEB_SCRAPING_AVAILABLE = False
    print("Web scraping libraries not available. Install with: pip install requests beautifulsoup4 lxml html2text")

# Advanced web content extraction
try:
    from newspaper import Article
    import newspaper
    NEWSPAPER_AVAILABLE = True
except ImportError:
    NEWSPAPER_AVAILABLE = False
    print("Newspaper3k not available. Install with: pip install newspaper3k")

# User agent spoofing
try:
    from fake_useragent import UserAgent
    FAKE_USERAGENT_AVAILABLE = True
except ImportError:
    FAKE_USERAGENT_AVAILABLE = False
    print("Fake UserAgent not available. Install with: pip install fake-useragent")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@dataclass
class ExtractionResult:
    """Data class for extraction results"""
    text: str
    confidence: float
    language: str
    content_type: str
    metadata: Dict[str, Any]
    extraction_time: float

class AdvancedTextExtractor:
    """
    Advanced text extraction system supporting multiple content types and languages
    """

    def __init__(self,
                 languages: List[str] = ['en', 'ch_sim'],
                 output_dir: str = "extracted_texts",
                 use_gpu: bool = False):
        """
        Initialize the text extractor

        Args:
            languages: List of language codes (ISO format)
            output_dir: Directory to save extracted text files
            use_gpu: Whether to use GPU acceleration (if available)
        """
        self.languages = languages
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.use_gpu = use_gpu

        # Language mapping
        self.lang_map = {
            'en': 'english',
            'ch_sim': 'chinese_simplified',
            'ch_tra': 'chinese_traditional',
            'zh': 'chinese_simplified',
            'ja': 'japanese',
            'ko': 'korean',
            'ar': 'arabic',
            'hi': 'hindi',
            'es': 'spanish',
            'fr': 'french',
            'de': 'german',
            'ru': 'russian'
        }

        # Initialize OCR engines
        self._init_ocr_engines()

        # Initialize document processors
        self._init_document_processors()

        # Initialize web scraping
        self._init_web_scraping()

        logger.info(f"TextExtractor initialized with languages: {languages}")

    def _init_ocr_engines(self):
        """Initialize OCR engines"""
        try:
            # EasyOCR for multilingual support
            self.easyocr_reader = easyocr.Reader(self.languages, gpu=self.use_gpu)

            # PaddleOCR for enhanced accuracy (especially for Chinese)
            if PADDLE_AVAILABLE:
                # PaddleOCR language mapping
                paddle_lang_map = {
                    'en': 'en',
                    'ch_sim': 'ch',
                    'ch_tra': 'chinese_cht',
                    'zh': 'ch',
                    'ja': 'japan',
                    'ko': 'korean',
                    'ar': 'ar',
                    'hi': 'hi',
                    'es': 'es',
                    'fr': 'french',
                    'de': 'german',
                    'ru': 'ru'
                }

                # Find primary language for PaddleOCR (it works best with single language)
                primary_lang = None
                for lang in self.languages:
                    if lang in paddle_lang_map:
                        primary_lang = paddle_lang_map[lang]
                        break

                if primary_lang:
                    try:
                        self.paddle_ocr = PaddleOCR(
                            use_angle_cls=True,
                            lang=primary_lang,
                            use_gpu=self.use_gpu,
                            show_log=False  # Reduce verbose output
                        )
                        logger.info(f"PaddleOCR initialized with language: {primary_lang}")
                    except Exception as paddle_error:
                        logger.warning(f"PaddleOCR initialization failed: {paddle_error}")
                        self.paddle_ocr = None
                else:
                    logger.info("No suitable language found for PaddleOCR")
                    self.paddle_ocr = None
            else:
                self.paddle_ocr = None

            # Configure Tesseract
            tesseract_langs = '+'.join([self.lang_map.get(lang, lang) for lang in self.languages])
            self.tesseract_config = f'--oem 3 --psm 6 -l {tesseract_langs}'

            logger.info("OCR engines initialized successfully")

        except Exception as e:
            logger.error(f"Error initializing OCR engines: {e}")
            # Don't raise the error, continue with available engines
            logger.warning("Some OCR engines failed to initialize, continuing with available ones")

    def _init_document_processors(self):
        """Initialize document processing tools"""
        self.supported_formats = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.doc': self._extract_from_docx,
            '.txt': self._extract_from_txt,
            '.jpg': self._extract_from_image,
            '.jpeg': self._extract_from_image,
            '.png': self._extract_from_image,
            '.bmp': self._extract_from_image,
            '.tiff': self._extract_from_image,
            '.webp': self._extract_from_image,
            '.xlsx': self._extract_from_excel,
            '.xls': self._extract_from_excel,
            '.csv': self._extract_from_csv
        }

    def _init_web_scraping(self):
        """Initialize web scraping capabilities"""
        if WEB_SCRAPING_AVAILABLE:
            # Setup session with headers
            self.session = requests.Session()

            # User agent rotation
            if FAKE_USERAGENT_AVAILABLE:
                try:
                    ua = UserAgent()
                    self.session.headers.update({
                        'User-Agent': ua.random
                    })
                except:
                    # Fallback user agent
                    self.session.headers.update({
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                    })
            else:
                self.session.headers.update({
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                })

            # HTML to text converter
            self.html_converter = html2text.HTML2Text()
            self.html_converter.ignore_links = False
            self.html_converter.ignore_images = False
            self.html_converter.body_width = 0  # Don't wrap lines

            logger.info("Web scraping initialized successfully")
        else:
            self.session = None
            self.html_converter = None
            logger.warning("Web scraping not available - install required packages")

    def extract_text(self, source: Union[str, Path]) -> ExtractionResult:
        """
        Extract text from various sources (files or URLs)

        Args:
            source: Path to file or URL to extract text from

        Returns:
            ExtractionResult object containing extracted text and metadata
        """
        start_time = datetime.now()
        source_str = str(source)

        # Check if source is a URL
        if self._is_url(source_str):
            return self._extract_from_url(source_str, start_time)

        # Handle as file path
        file_path = Path(source_str)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = file_path.suffix.lower()

        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        try:
            # Extract text using appropriate method
            extractor_func = self.supported_formats[file_extension]
            text, confidence, detected_lang, metadata = extractor_func(file_path)

            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()

            # Create result object
            result = ExtractionResult(
                text=text,
                confidence=confidence,
                language=detected_lang,
                content_type=file_extension[1:],  # Remove the dot
                metadata=metadata,
                extraction_time=processing_time
            )

            # Save to text file
            output_file = self._save_extracted_text(source, result)
            result.metadata['output_file'] = str(output_file)

            logger.info(f"Successfully extracted text from {file_path.name} in {processing_time:.2f}s")
            return result

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            logger.error(traceback.format_exc())
            raise

    def _is_url(self, source: str) -> bool:
        """Check if source is a valid URL"""
        try:
            result = urlparse(source)
            return all([result.scheme, result.netloc]) and result.scheme in ['http', 'https']
        except:
            return False

    def _extract_from_url(self, url: str, start_time: datetime) -> ExtractionResult:
        """Extract text from URL"""
        if not WEB_SCRAPING_AVAILABLE:
            raise ValueError("Web scraping not available. Install required packages: pip install requests beautifulsoup4 lxml html2text newspaper3k")

        try:
            logger.info(f"Processing URL: {url}")

            # Determine URL content type
            url_type, file_extension = self._analyze_url(url)

            if url_type == "direct_file":
                # Direct file URL (PDF, image, etc.)
                return self._extract_from_url_file(url, file_extension, start_time)
            elif url_type == "webpage":
                # Web page content
                return self._extract_from_webpage(url, start_time)
            else:
                raise ValueError(f"Unsupported URL type: {url_type}")

        except Exception as e:
            logger.error(f"Error extracting from URL {url}: {e}")
            raise

    def _analyze_url(self, url: str) -> tuple:
        """Analyze URL to determine content type"""
        parsed = urlparse(url)
        path = parsed.path.lower()

        # Check for direct file URLs
        file_extensions = ['.pdf', '.doc', '.docx', '.txt', '.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp', '.xlsx', '.xls', '.csv']

        for ext in file_extensions:
            if path.endswith(ext):
                return "direct_file", ext

        # Check content type via HEAD request
        try:
            response = self.session.head(url, timeout=10, allow_redirects=True)
            content_type = response.headers.get('content-type', '').lower()

            if 'application/pdf' in content_type:
                return "direct_file", ".pdf"
            elif 'image/' in content_type:
                # Determine image extension from content type
                if 'jpeg' in content_type:
                    return "direct_file", ".jpg"
                elif 'png' in content_type:
                    return "direct_file", ".png"
                else:
                    return "direct_file", ".jpg"  # Default
            elif any(x in content_type for x in ['excel', 'spreadsheet']):
                return "direct_file", ".xlsx"
            elif 'text/csv' in content_type:
                return "direct_file", ".csv"
            elif any(x in content_type for x in ['msword', 'wordprocessingml']):
                return "direct_file", ".docx"
        except:
            pass

        # Default to webpage
        return "webpage", ".html"

    def _extract_from_url_file(self, url: str, file_extension: str, start_time: datetime) -> ExtractionResult:
        """Extract text from direct file URL"""
        import tempfile

        try:
            # Download file to temporary location
            with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
                temp_path = temp_file.name

                response = self.session.get(url, timeout=30, stream=True)
                response.raise_for_status()

                # Download with progress
                total_size = int(response.headers.get('content-length', 0))
                downloaded = 0

                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        temp_file.write(chunk)
                        downloaded += len(chunk)

                        if total_size > 0:
                            progress = (downloaded / total_size) * 100
                            if downloaded % (1024 * 1024) == 0:  # Log every MB
                                logger.info(f"Downloaded {progress:.1f}% ({downloaded / 1024 / 1024:.1f} MB)")

            # Process the downloaded file
            if file_extension in self.supported_formats:
                extractor_func = self.supported_formats[file_extension]
                text, confidence, detected_lang, metadata = extractor_func(Path(temp_path))

                # Add URL metadata
                metadata['source_url'] = url
                metadata['download_size'] = downloaded
                metadata['file_type'] = 'url_download'
            else:
                raise ValueError(f"Unsupported file extension from URL: {file_extension}")

            # Clean up temporary file
            try:
                os.unlink(temp_path)
            except:
                pass

            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()

            # Create result
            result = ExtractionResult(
                text=text,
                confidence=confidence,
                language=detected_lang,
                content_type=f"url_{file_extension[1:]}",
                metadata=metadata,
                extraction_time=processing_time
            )

            # Save to text file
            output_file = self._save_extracted_text(url, result)
            result.metadata['output_file'] = str(output_file)

            logger.info(f"Successfully extracted text from URL file in {processing_time:.2f}s")
            return result

        except Exception as e:
            logger.error(f"Error processing URL file {url}: {e}")
            raise

    def _extract_from_webpage(self, url: str, start_time: datetime) -> ExtractionResult:
        """Extract text from web page"""
        try:
            # Try newspaper3k first for article extraction
            if NEWSPAPER_AVAILABLE:
                try:
                    article_text, article_confidence, article_metadata = self._extract_with_newspaper(url)
                    if article_text.strip():
                        processing_time = (datetime.now() - start_time).total_seconds()

                        result = ExtractionResult(
                            text=article_text,
                            confidence=article_confidence,
                            language=self._detect_language(article_text),
                            content_type="webpage_article",
                            metadata=article_metadata,
                            extraction_time=processing_time
                        )

                        # Save to text file
                        output_file = self._save_extracted_text(url, result)
                        result.metadata['output_file'] = str(output_file)

                        logger.info(f"Successfully extracted article from {url} in {processing_time:.2f}s")
                        return result
                except Exception as e:
                    logger.warning(f"Newspaper3k extraction failed: {e}")

            # Fallback to BeautifulSoup extraction
            text, confidence, metadata = self._extract_with_beautifulsoup(url)

            processing_time = (datetime.now() - start_time).total_seconds()

            result = ExtractionResult(
                text=text,
                confidence=confidence,
                language=self._detect_language(text),
                content_type="webpage_html",
                metadata=metadata,
                extraction_time=processing_time
            )

            # Save to text file
            output_file = self._save_extracted_text(url, result)
            result.metadata['output_file'] = str(output_file)

            logger.info(f"Successfully extracted webpage content from {url} in {processing_time:.2f}s")
            return result

        except Exception as e:
            logger.error(f"Error extracting from webpage {url}: {e}")
            raise

    def _extract_with_newspaper(self, url: str) -> tuple:
        """Extract article content using newspaper3k"""
        try:
            article = Article(url)
            article.download()
            article.parse()

            # Get main content
            text_parts = []

            if article.title:
                text_parts.append(f"Title: {article.title}")

            if article.text:
                text_parts.append(f"Content:\n{article.text}")

            if article.summary:
                text_parts.append(f"Summary:\n{article.summary}")

            full_text = '\n\n'.join(text_parts)

            metadata = {
                'source_url': url,
                'title': article.title or '',
                'authors': article.authors,
                'publish_date': str(article.publish_date) if article.publish_date else '',
                'extraction_method': 'newspaper3k',
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }

            confidence = 0.9 if len(full_text.strip()) > 100 else 0.5

            return full_text, confidence, metadata

        except Exception as e:
            logger.warning(f"Newspaper3k extraction failed: {e}")
            raise

    def _extract_with_beautifulsoup(self, url: str) -> tuple:
        """Extract content using BeautifulSoup"""
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()

            # Parse HTML
            soup = BeautifulSoup(response.content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
                script.decompose()

            # Extract title
            title = soup.find('title')
            title_text = title.get_text().strip() if title else ''

            # Try to find main content areas
            main_content = None

            # Look for common content containers
            content_selectors = [
                'main', 'article', '[role="main"]',
                '.content', '.main-content', '.article-content',
                '#content', '#main-content', '#article-content',
                '.post-content', '.entry-content', '.page-content'
            ]

            for selector in content_selectors:
                main_content = soup.select_one(selector)
                if main_content:
                    break

            # If no main content found, use body
            if not main_content:
                main_content = soup.find('body')

            if not main_content:
                main_content = soup

            # Convert to text
            if self.html_converter:
                # Use html2text for better formatting
                html_content = str(main_content)
                text_content = self.html_converter.handle(html_content)
            else:
                # Fallback to simple text extraction
                text_content = main_content.get_text()

            # Clean up text
            lines = text_content.split('\n')
            cleaned_lines = []

            for line in lines:
                line = line.strip()
                if line and len(line) > 3:  # Skip very short lines
                    cleaned_lines.append(line)

            final_text = '\n'.join(cleaned_lines)

            # Add title if available
            if title_text:
                final_text = f"Title: {title_text}\n\n{final_text}"

            metadata = {
                'source_url': url,
                'title': title_text,
                'extraction_method': 'beautifulsoup',
                'content_length': len(response.content),
                'word_count': len(final_text.split()),
                'char_count': len(final_text),
                'response_status': response.status_code
            }

            # Determine confidence based on content quality
            if len(final_text.strip()) > 500:
                confidence = 0.8
            elif len(final_text.strip()) > 100:
                confidence = 0.6
            else:
                confidence = 0.3

            return final_text, confidence, metadata

        except Exception as e:
            logger.error(f"BeautifulSoup extraction failed: {e}")
            raise

    def _extract_from_pdf(self, file_path: Path) -> tuple:
        """Extract text from PDF files"""
        text_parts = []
        confidence = 0.0
        detected_lang = 'en'
        metadata = {'pages': 0, 'tables': [], 'images': 0}

        try:
            # Try text extraction first (for text-based PDFs)
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['pages'] = len(pdf_reader.pages)

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")

            # If no text extracted, use OCR
            if not any(text_parts):
                logger.info(f"No text found in PDF {file_path.name}, using OCR...")
                text_parts, confidence, detected_lang = self._ocr_pdf_pages(file_path)
            else:
                confidence = 0.95  # High confidence for text-based PDFs

            # Extract tables using camelot
            try:
                tables = camelot.read_pdf(str(file_path), pages='all')
                for i, table in enumerate(tables):
                    if len(table.df) > 0:
                        table_text = f"\n--- Table {i+1} ---\n"
                        table_text += table.df.to_string(index=False)
                        text_parts.append(table_text)
                        metadata['tables'].append({
                            'table_id': i+1,
                            'rows': len(table.df),
                            'columns': len(table.df.columns),
                            'accuracy': table.accuracy
                        })
            except Exception as e:
                logger.warning(f"Could not extract tables from PDF: {e}")

            # Extract images and apply OCR
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        images = page.images
                        metadata['images'] += len(images)

                        for img_idx, img in enumerate(images):
                            # Extract image and apply OCR if it's text-heavy
                            try:
                                img_obj = page.crop(img['bbox']).to_image()
                                img_text = self._extract_text_from_image_object(img_obj.original)
                                if img_text.strip():
                                    text_parts.append(f"\n--- Image {img_idx+1} on Page {page_num+1} ---\n{img_text}")
                            except Exception as img_e:
                                logger.warning(f"Could not process image {img_idx+1} on page {page_num+1}: {img_e}")

            except Exception as e:
                logger.warning(f"Could not extract images from PDF: {e}")

        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise

        final_text = '\n\n'.join(text_parts) if text_parts else ""
        return final_text, confidence, detected_lang, metadata

    def _ocr_pdf_pages(self, file_path: Path) -> tuple:
        """Apply OCR to PDF pages"""
        try:
            import fitz  # PyMuPDF
            import io
        except ImportError:
            logger.warning("PyMuPDF not available for PDF OCR. Install with: pip install PyMuPDF")
            return [], 0.0, 'en'

        text_parts = []
        confidences = []
        detected_langs = []

        try:
            doc = fitz.open(file_path)

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")

                # Convert to PIL Image
                img = Image.open(io.BytesIO(img_data))

                # Apply OCR
                page_text, conf, lang = self._extract_text_from_image_object(img)

                if page_text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    confidences.append(conf)
                    detected_langs.append(lang)

            doc.close()

        except Exception as e:
            logger.error(f"Error during PDF OCR: {e}")
            return [], 0.0, 'en'

        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
        primary_lang = max(set(detected_langs), key=detected_langs.count) if detected_langs else 'en'

        return text_parts, avg_confidence, primary_lang

    def _extract_from_docx(self, file_path: Path) -> tuple:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            tables_data = []
            for table_idx, table in enumerate(doc.tables):
                table_text = f"\n--- Table {table_idx + 1} ---\n"
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_data)

                if table_rows:
                    df = pd.DataFrame(table_rows[1:], columns=table_rows[0] if table_rows else [])
                    table_text += df.to_string(index=False)
                    text_parts.append(table_text)
                    tables_data.append({
                        'table_id': table_idx + 1,
                        'rows': len(table_rows),
                        'columns': len(table_rows[0]) if table_rows else 0
                    })

            full_text = '\n\n'.join(text_parts)

            # Simple language detection based on character sets
            detected_lang = self._detect_language(full_text)

            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': tables_data,
                'word_count': len(full_text.split())
            }

            return full_text, 0.95, detected_lang, metadata

        except Exception as e:
            logger.error(f"Error extracting from DOCX {file_path}: {e}")
            raise

    def _extract_from_txt(self, file_path: Path) -> tuple:
        """Extract text from TXT files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'cp1252', 'latin1', 'gbk', 'big5']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError(f"Could not decode {file_path} with any supported encoding")

            detected_lang = self._detect_language(text)

            metadata = {
                'encoding': encoding,
                'line_count': len(text.splitlines()),
                'word_count': len(text.split())
            }

            return text, 1.0, detected_lang, metadata

        except Exception as e:
            logger.error(f"Error extracting from TXT {file_path}: {e}")
            raise

    def _extract_from_image(self, file_path: Path) -> tuple:
        """Extract text from image files"""
        try:
            img = Image.open(file_path)
            return self._extract_text_from_image_object(img)
        except Exception as e:
            logger.error(f"Error extracting from image {file_path}: {e}")
            raise

    def _extract_text_from_image_object(self, img: Image.Image) -> tuple:
        """Extract text from PIL Image object using multiple OCR engines"""
        # Convert to RGB if necessary
        if img.mode != 'RGB':
            img = img.convert('RGB')

        # Convert to numpy array for OpenCV processing
        img_np = np.array(img)

        # Preprocess image for better OCR
        img_processed = self._preprocess_image(img_np)

        results = []

        # Try EasyOCR first (most reliable for multilingual)
        try:
            easyocr_result = self.easyocr_reader.readtext(img_processed)
            easyocr_text = ' '.join([detection[1] for detection in easyocr_result])
            easyocr_conf = np.mean([detection[2] for detection in easyocr_result]) if easyocr_result else 0

            if easyocr_text.strip():
                results.append({
                    'text': easyocr_text,
                    'confidence': easyocr_conf,
                    'engine': 'EasyOCR'
                })
        except Exception as e:
            logger.warning(f"EasyOCR failed: {e}")

        # Try Tesseract (fallback)
        try:
            tesseract_text = pytesseract.image_to_string(img_processed, config=self.tesseract_config)
            if tesseract_text.strip():
                # Get confidence from Tesseract
                try:
                    tessdata = pytesseract.image_to_data(img_processed, config=self.tesseract_config, output_type=pytesseract.Output.DICT)
                    confidences = [int(conf) for conf in tessdata['conf'] if int(conf) > 0]
                    tesseract_conf = np.mean(confidences) / 100.0 if confidences else 0
                except:
                    tesseract_conf = 0.7  # Default confidence

                results.append({
                    'text': tesseract_text,
                    'confidence': tesseract_conf,
                    'engine': 'Tesseract'
                })
        except Exception as e:
            logger.warning(f"Tesseract failed: {e}")

        # Try PaddleOCR (if available and initialized)
        if self.paddle_ocr:
            try:
                paddle_result = self.paddle_ocr.ocr(img_processed, cls=True)
                if paddle_result and paddle_result[0]:
                    paddle_texts = []
                    paddle_confs = []

                    for line in paddle_result[0]:
                        if line and len(line) >= 2:
                            text_info = line[1]
                            if isinstance(text_info, (list, tuple)) and len(text_info) >= 2:
                                paddle_texts.append(text_info[0])
                                paddle_confs.append(text_info[1])

                    if paddle_texts:
                        paddle_text = ' '.join(paddle_texts)
                        paddle_conf = np.mean(paddle_confs)

                        results.append({
                            'text': paddle_text,
                            'confidence': paddle_conf,
                            'engine': 'PaddleOCR'
                        })
            except Exception as e:
                logger.warning(f"PaddleOCR failed: {e}")

        # Select best result
        if not results:
            logger.warning("All OCR engines failed, returning empty result")
            return "", 0.0, 'en'

        # Choose result with highest confidence, but prefer non-empty results
        valid_results = [r for r in results if len(r['text'].strip()) > 5]  # At least 5 characters
        if valid_results:
            best_result = max(valid_results, key=lambda x: x['confidence'])
        else:
            best_result = max(results, key=lambda x: x['confidence'])

        detected_lang = self._detect_language(best_result['text'])

        logger.info(f"Best OCR result from {best_result['engine']} with confidence {best_result['confidence']:.2f}")

        return best_result['text'], best_result['confidence'], detected_lang

    def _preprocess_image(self, img: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR results"""
        # Convert to grayscale
        if len(img.shape) == 3:
            gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
        else:
            gray = img

        # Apply Gaussian blur to reduce noise
        blurred = cv2.GaussianBlur(gray, (5, 5), 0)

        # Apply threshold to get binary image
        _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # Morphological operations to clean up the image
        kernel = np.ones((1, 1), np.uint8)
        processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
        processed = cv2.morphologyEx(processed, cv2.MORPH_OPEN, kernel)

        return processed

    def _extract_from_excel(self, file_path: Path) -> tuple:
        """Extract text from Excel files"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_parts = []

            metadata = {
                'sheets': [],
                'total_rows': 0,
                'total_columns': 0
            }

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                if not df.empty:
                    sheet_text = f"\n--- Sheet: {sheet_name} ---\n"
                    sheet_text += df.to_string(index=False)
                    text_parts.append(sheet_text)

                    metadata['sheets'].append({
                        'name': sheet_name,
                        'rows': len(df),
                        'columns': len(df.columns)
                    })
                    metadata['total_rows'] += len(df)
                    metadata['total_columns'] += len(df.columns)

            full_text = '\n\n'.join(text_parts)
            detected_lang = self._detect_language(full_text)

            return full_text, 0.95, detected_lang, metadata

        except Exception as e:
            logger.error(f"Error extracting from Excel {file_path}: {e}")
            raise

    def _extract_from_csv(self, file_path: Path) -> tuple:
        """Extract text from CSV files"""
        try:
            # Try different encodings and separators
            encodings = ['utf-8', 'cp1252', 'latin1', 'gbk']
            separators = [',', ';', '\t']

            df = None
            used_encoding = 'utf-8'
            used_separator = ','

            for encoding in encodings:
                for sep in separators:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding, sep=sep)
                        if len(df.columns) > 1:  # Valid CSV should have multiple columns
                            used_encoding = encoding
                            used_separator = sep
                            break
                    except:
                        continue
                if df is not None and len(df.columns) > 1:
                    break

            if df is None or len(df.columns) <= 1:
                raise ValueError(f"Could not properly parse CSV file {file_path}")

            text = df.to_string(index=False)
            detected_lang = self._detect_language(text)

            metadata = {
                'encoding': used_encoding,
                'separator': used_separator,
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns)
            }

            return text, 0.95, detected_lang, metadata

        except Exception as e:
            logger.error(f"Error extracting from CSV {file_path}: {e}")
            raise

    def _detect_language(self, text: str) -> str:
        """Simple language detection based on character sets"""
        if not text:
            return 'en'

        # Count character types
        chinese_chars = len([c for c in text if '\u4e00' <= c <= '\u9fff'])
        japanese_chars = len([c for c in text if '\u3040' <= c <= '\u309f' or '\u30a0' <= c <= '\u30ff'])
        korean_chars = len([c for c in text if '\uac00' <= c <= '\ud7af'])
        arabic_chars = len([c for c in text if '\u0600' <= c <= '\u06ff'])

        total_chars = len([c for c in text if c.isalpha()])

        if total_chars == 0:
            return 'en'

        # Determine language based on character distribution
        if chinese_chars / total_chars > 0.3:
            return 'zh'
        elif japanese_chars / total_chars > 0.3:
            return 'ja'
        elif korean_chars / total_chars > 0.3:
            return 'ko'
        elif arabic_chars / total_chars > 0.3:
            return 'ar'
        else:
            return 'en'

    def _save_extracted_text(self, source: Union[Path, str], result: ExtractionResult) -> Path:
        """Save extracted text to a file"""
        # Create output filename
        if isinstance(source, (str,)) and self._is_url(source):
            # Handle URL
            parsed = urlparse(source)
            domain = parsed.netloc.replace('.', '_')
            path_part = parsed.path.replace('/', '_').replace('\\', '_')
            base_name = f"url_{domain}_{path_part}"[:50]  # Limit length
        else:
            # Handle file path
            base_name = Path(source).stem

        # Clean filename
        base_name = "".join(c for c in base_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
        if not base_name:
            base_name = "extracted_content"

        output_file = self.output_dir / f"{base_name}_extracted_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

        # Prepare content
        content = []
        content.append(f"=== TEXT EXTRACTION REPORT ===")
        content.append(f"Source: {source}")
        content.append(f"Extraction Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append(f"Content Type: {result.content_type}")
        content.append(f"Detected Language: {result.language}")
        content.append(f"Confidence Score: {result.confidence:.2f}")
        content.append(f"Processing Time: {result.extraction_time:.2f} seconds")
        content.append(f"Metadata: {json.dumps(result.metadata, indent=2)}")
        content.append("=" * 50)
        content.append("")
        content.append("=== EXTRACTED TEXT ===")
        content.append(result.text)

        # Save to file
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write('\n'.join(content))

            logger.info(f"Extracted text saved to: {output_file}")
            return output_file

        except Exception as e:
            logger.error(f"Error saving extracted text: {e}")
            raise

    def batch_extract(self, sources: List[Union[str, Path]], file_pattern: str = "*") -> List[ExtractionResult]:
        """
        Extract text from multiple sources (files or URLs)

        Args:
            sources: List of file paths, directory paths, or URLs
            file_pattern: File pattern to match for directories (e.g., "*.pdf", "*.jpg")

        Returns:
            List of ExtractionResult objects
        """
        all_sources = []

        # Process each source
        for source in sources:
            source_str = str(source)

            if self._is_url(source_str):
                # Add URL directly
                all_sources.append(source_str)
            else:
                # Handle as file/directory path
                source_path = Path(source_str)

                if not source_path.exists():
                    logger.warning(f"Path not found: {source_path}")
                    continue

                if source_path.is_file():
                    # Single file
                    if source_path.suffix.lower() in self.supported_formats:
                        all_sources.append(source_path)
                    else:
                        logger.warning(f"Unsupported file format: {source_path}")
                elif source_path.is_dir():
                    # Directory - find matching files
                    matching_files = list(source_path.glob(file_pattern))
                    for file_path in matching_files:
                        if file_path.suffix.lower() in self.supported_formats:
                            all_sources.append(file_path)

        if not all_sources:
            logger.warning("No valid sources found for processing")
            return []

        results = []

        for i, source in enumerate(all_sources, 1):
            source_name = str(source)
            if len(source_name) > 50:
                display_name = source_name[:47] + "..."
            else:
                display_name = source_name

            try:
                logger.info(f"Processing {i}/{len(all_sources)}: {display_name}")
                result = self.extract_text(source)
                results.append(result)
                logger.info(f"  ✅ Success - Confidence: {result.confidence:.2f}, Language: {result.language}")
            except Exception as e:
                logger.error(f"  ❌ Failed to process {display_name}: {e}")
                continue

        logger.info(f"Batch processing completed. {len(results)} sources processed successfully.")
        return results

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats and URL types"""
        formats = list(self.supported_formats.keys())
        formats.extend(['URLs (http/https)', 'Web pages', 'Online PDFs', 'Online images'])
        return formats

    def create_extraction_summary(self, results: List[ExtractionResult]) -> str:
        """Create a summary of extraction results"""
        if not results:
            return "No extraction results to summarize."

        summary = []
        summary.append("=== EXTRACTION SUMMARY ===")
        summary.append(f"Total Files Processed: {len(results)}")
        summary.append(f"Total Processing Time: {sum(r.extraction_time for r in results):.2f} seconds")
        summary.append("")

        # Group by content type
        by_type = {}
        for result in results:
            if result.content_type not in by_type:
                by_type[result.content_type] = []
            by_type[result.content_type].append(result)

        for content_type, type_results in by_type.items():
            summary.append(f"{content_type.upper()} Files: {len(type_results)}")
            avg_confidence = sum(r.confidence for r in type_results) / len(type_results)
            summary.append(f"  Average Confidence: {avg_confidence:.2f}")

            # Language distribution
            langs = [r.language for r in type_results]
            lang_counts = {lang: langs.count(lang) for lang in set(langs)}
            summary.append(f"  Languages: {dict(sorted(lang_counts.items()))}")
            summary.append("")

        # Save summary
        summary_file = self.output_dir / f"extraction_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        with open(summary_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(summary))

        logger.info(f"Extraction summary saved to: {summary_file}")
        return '\n'.join(summary)

def main():
    """Main function with interactive user input"""
    print("🚀 Advanced Multilingual Text Extraction System")
    print("=" * 60)

    # Check web scraping availability
    if not WEB_SCRAPING_AVAILABLE:
        print("⚠️ Web scraping not available. To process URLs, install:")
        print("pip install requests beautifulsoup4 lxml html2text newspaper3k fake-useragent")
        print()

    # Get user preferences
    print("\n📝 System Configuration:")

    # Languages
    print("Available languages: en (English), ch_sim (Chinese Simplified), ch_tra (Chinese Traditional)")
    print("ja (Japanese), ko (Korean), ar (Arabic), hi (Hindi), es (Spanish), fr (French), de (German), ru (Russian)")
    lang_input = input("Enter languages (comma-separated, default: en,ch_sim): ").strip()
    if lang_input:
        languages = [lang.strip() for lang in lang_input.split(',')]
    else:
        languages = ["en", "ch_sim"]

    # Output directory
    output_dir = input("Output directory (default: extracted_texts): ").strip()
    if not output_dir:
        output_dir = "extracted_texts"

    # GPU usage
    gpu_input = input("Use GPU acceleration if available? (y/n, default: n): ").strip().lower()
    use_gpu = gpu_input in ['y', 'yes', '1', 'true']

    # Initialize extractor
    print(f"\n⚙️ Initializing extractor with languages: {languages}")
    extractor = AdvancedTextExtractor(
        languages=languages,
        output_dir=output_dir,
        use_gpu=use_gpu
    )

    print(f"📁 Supported formats: {', '.join(extractor.get_supported_formats())}")

    try:
        # Processing mode
        print("\n📂 Processing Mode:")
        print("1. Single file processing")
        print("2. Single URL processing")
        print("3. Batch processing (files in folder)")
        print("4. Multiple files/URLs")
        print("5. URL list from file")

        mode = input("Choose mode (1-5, default: 1): ").strip()

        if mode == "2":
            # Single URL processing
            url = input("\n🌐 Enter URL: ").strip()
            if not url:
                print("❌ No URL provided!")
                return

            print(f"\n🔄 Processing: {url}")
            result = extractor.extract_text(url)

            print("\n✅ Text extracted successfully!")
            print(f"📄 Output file: {result.metadata.get('output_file', 'N/A')}")
            print(f"🎯 Confidence: {result.confidence:.2f}")
            print(f"🌐 Language: {result.language}")
            print(f"⏱️ Processing time: {result.extraction_time:.2f} seconds")
            print(f"📊 Content type: {result.content_type}")

            # Show preview
            preview_input = input("\n👀 Show text preview? (y/n, default: y): ").strip().lower()
            if preview_input != 'n':
                preview_text = result.text[:500] + "..." if len(result.text) > 500 else result.text
                print(f"\n📖 Text Preview:\n{'-'*40}\n{preview_text}\n{'-'*40}")

        elif mode == "3":
            # Batch processing (folder)
            folder_path = input("\n📂 Enter folder path to process: ").strip()
            if not folder_path:
                print("❌ No folder path provided!")
                return

            pattern = input("File pattern (default: *, examples: *.pdf, *.jpg): ").strip()
            if not pattern:
                pattern = "*"

            print(f"\n🔄 Processing all files in '{folder_path}' matching '{pattern}'...")
            results = extractor.batch_extract([folder_path], pattern)

            if results:
                summary_input = input("\n📊 Generate processing summary? (y/n, default: y): ").strip().lower()
                if summary_input != 'n':
                    summary = extractor.create_extraction_summary(results)
                    print("\n" + "="*60)
                    print(summary)

                print(f"\n✅ Batch processing completed! {len(results)} files processed.")
            else:
                print("❌ No files were processed!")

        elif mode == "4":
            # Multiple files/URLs
            print("\n📄 Enter file paths or URLs (one per line, empty line to finish):")
            sources = []
            while True:
                source = input("File path or URL: ").strip()
                if not source:
                    break
                sources.append(source)

            if not sources:
                print("❌ No sources provided!")
                return

            print(f"\n🔄 Processing {len(sources)} sources...")
            results = extractor.batch_extract(sources)

            if results:
                summary_input = input("\n📊 Generate processing summary? (y/n, default: y): ").strip().lower()
                if summary_input != 'n':
                    summary = extractor.create_extraction_summary(results)
                    print("\n" + "="*60)
                    print(summary)

                print(f"\n✅ Multiple source processing completed! {len(results)} sources processed successfully.")
            else:
                print("❌ No sources were processed successfully!")

        elif mode == "5":
            # URL list from file
            url_file = input("\n📄 Enter path to file containing URLs (one per line): ").strip()
            if not url_file:
                print("❌ No file path provided!")
                return

            try:
                with open(url_file, 'r', encoding='utf-8') as f:
                    urls = [line.strip() for line in f if line.strip() and not line.startswith('#')]

                if not urls:
                    print("❌ No valid URLs found in file!")
                    return

                print(f"\n🔄 Processing {len(urls)} URLs from file...")
                results = extractor.batch_extract(urls)

                if results:
                    summary_input = input("\n📊 Generate processing summary? (y/n, default: y): ").strip().lower()
                    if summary_input != 'n':
                        summary = extractor.create_extraction_summary(results)
                        print("\n" + "="*60)
                        print(summary)

                    print(f"\n✅ URL batch processing completed! {len(results)} URLs processed successfully.")
                else:
                    print("❌ No URLs were processed successfully!")

            except Exception as e:
                print(f"❌ Error reading URL file: {e}")
                return

        else:
            # Single file processing
            file_path = input("\n📄 Enter file path: ").strip()
            if not file_path:
                print("❌ No file path provided!")
                return

            print(f"\n🔄 Processing: {Path(file_path).name}")
            result = extractor.extract_text(file_path)

            print("\n✅ Text extracted successfully!")
            print(f"📄 Output file: {result.metadata.get('output_file', 'N/A')}")
            print(f"🎯 Confidence: {result.confidence:.2f}")
            print(f"🌐 Language: {result.language}")
            print(f"⏱️ Processing time: {result.extraction_time:.2f} seconds")
            print(f"📊 Content type: {result.content_type}")

            # Show preview
            preview_input = input("\n👀 Show text preview? (y/n, default: y): ").strip().lower()
            if preview_input != 'n':
                preview_text = result.text[:500] + "..." if len(result.text) > 500 else result.text
                print(f"\n📖 Text Preview:\n{'-'*40}\n{preview_text}\n{'-'*40}")

    except Exception as e:
        logger.error(f"Extraction failed: {e}")
        print(f"❌ Error: {e}")
        return

    print(f"\n🎉 Processing completed! Check the '{output_dir}' folder for extracted text files.")

if __name__ == "__main__":
    main()